
from pathlib import Path
from typing import Dict, List, Tuple, Set
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

def normalize_key(s: str) -> str:
    if s is None:
        return ""
    s = s.strip()
    s = re.sub(r'^\[|\]$', '', s)
    s = re.sub(r'^\[\[?|\]?\]$', '', s)
    s = re.sub(r'^\{\{?|\}?\}$', '', s)
    s = re.sub(r'\s+', ' ', s)
    return s.upper()

def iter_paragraphs_and_cells(doc: Document):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for p in ncell.paragraphs:
                                yield p

def get_highlighted_runs(paragraph: Paragraph):
    runs = []
    for run in paragraph.runs:
        try:
            if run.font.highlight_color is not None:
                runs.append(run)
        except Exception:
            pass
    return runs

def extract_highlighted_labels(doc: Document) -> List[str]:
    labels = []
    seen = set()
    for p in iter_paragraphs_and_cells(doc):
        for run in get_highlighted_runs(p):
            text = run.text.strip()
            if not text:
                continue
            key = normalize_key(text)
            if key and key not in seen:
                seen.add(key)
                labels.append(text)
    return labels

def build_mapping_from_data(data: Dict[str, str], found_labels: List[str]) -> Dict[str, str]:
    norm_data = {normalize_key(k): ("" if v is None else str(v)) for k, v in data.items()}
    mapping = dict(norm_data)
    for lbl in found_labels:
        nk = normalize_key(lbl)
        if nk in norm_data:
            mapping[nk] = norm_data[nk]
    return mapping

def fill_document_bytes(template_bytes: bytes, data: Dict[str, str]) -> bytes:
    from io import BytesIO
    bio = BytesIO(template_bytes)
    doc = Document(bio)

    labels = extract_highlighted_labels(doc)
    mapping = build_mapping_from_data(data, labels)

    # Replacement pass
    replacements = 0
    for p in iter_paragraphs_and_cells(doc):
        # highlighted runs fully replaced
        for run in get_highlighted_runs(p):
            raw = run.text.strip()
            nk = normalize_key(raw)
            if nk and nk in mapping:
                if run.text != mapping[nk]:
                    run.text = mapping[nk]
                    run.font.highlight_color = None
                    replacements += 1
        # token-level replacements
        for run in p.runs:
            before = run.text
            after = before
            for nk, val in mapping.items():
                for token in (nk, f'[{nk}]', f'[[{nk}]]', f'{{{{{nk}}}}}', f'__{nk}__'):
                    if token in after:
                        after = after.replace(token, val)
            if after != before:
                run.text = after
                replacements += 1

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
