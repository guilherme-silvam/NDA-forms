from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from werkzeug.utils import secure_filename
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Cm
import io

app = Flask(__name__)
app.secret_key = "change-this-key"

# Placeholders suportados (devem existir no template.docx)
# Ajustado para coletar apenas: Cliente, Endereço, CNPJ, Contrato e Data
# Use estes placeholders no seu template: {{CLIENTE}}, {{ENDERECO}}, {{CNPJ}}, {{CONTRATO}}, {{DATA}}
FIELDS = [
    "CLIENTE",        # Nome do cliente
    "ENDERECO",       # Endereço do cliente
    "CNPJ",           # CNPJ do cliente
    "CONTRATO",       # Número/identificação do contrato
    "DATA"            # Data (formato livre)
]

def replace_in_paragraph(paragraph, mapping: dict):
    """Troca {{CHAVE}} -> valor no parágrafo inteiro (recria o run para garantir a troca)."""
    if not paragraph.text:
        return
    text = paragraph.text
    changed = False
    for k, v in mapping.items():
        ph = "{{" + k + "}}"
        if ph in text:
            text = text.replace(ph, str(v))
            changed = True
    if changed:
        for r in paragraph.runs:
            r.text = ""
        paragraph.add_run(text)

def replace_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, mapping)
            for nested in cell.tables:
                replace_in_table(nested, mapping)

def insert_logo_placeholder(doc: Document, logo_path: Path, marker: str = "[LOGO]"):
    """
    Procura o marcador textual [LOGO] no documento e substitui por uma imagem.
    Busca em corpo, tabelas, cabeçalho e rodapé.
    """
    def handle_paragraph(p):
        if marker in (p.text or ""):
            for r in p.runs:
                r.text = ""
            run = p.add_run()
            try:
                run.add_picture(str(logo_path), width=Cm(5.0))
            except Exception:
                run.add_picture(str(logo_path))
            return True
        return False

    def handle_table(t):
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if handle_paragraph(p):
                        return True
                for nested in cell.tables:
                    if handle_table(nested):
                        return True
        return False

    for p in doc.paragraphs:
        if handle_paragraph(p):
            return True
    for t in doc.tables:
        if handle_table(t):
            return True
    for sec in doc.sections:
        if sec.header:
            for p in sec.header.paragraphs:
                if handle_paragraph(p):
                    return True
            for t in sec.header.tables:
                if handle_table(t):
                    return True
        if sec.footer:
            for p in sec.footer.paragraphs:
                if handle_paragraph(p):
                    return True
            for t in sec.footer.tables:
                if handle_table(t):
                    return True
    return False

def fill_docx(template_path: Path, data: dict, logo_path: Path | None = None) -> io.BytesIO:
    """Carrega template, substitui placeholders e insere o logo (se houver)."""
    doc = Document(template_path)

    # Trocas no corpo
    for p in doc.paragraphs:
        replace_in_paragraph(p, data)
    for t in doc.tables:
        replace_in_table(t, data)

    # Cabeçalho/rodapé
    for sec in doc.sections:
        if sec.header:
            for p in sec.header.paragraphs:
                replace_in_paragraph(p, data)
            for t in sec.header.tables:
                replace_in_table(t, data)
        if sec.footer:
            for p in sec.footer.paragraphs:
                replace_in_paragraph(p, data)
            for t in sec.footer.tables:
                replace_in_table(t, data)

    # Marca de logo
    if logo_path is not None:
        insert_logo_placeholder(doc, logo_path, marker="[LOGO]")

    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream

@app.route("/", methods=["GET"])
def index():
    # Seu form.html deve conter inputs com os names:
    # CLIENTE, ENDERECO, CNPJ, CONTRATO, DATA
    return render_template("form.html")

@app.route("/gerar", methods=["POST"])
def gerar():
    # Coleta dados do formulário
    data = {k: (request.form.get(k) or "").strip() for k in FIELDS}

    # Uploads (template e logo)
    uploads = Path("uploads")
    uploads.mkdir(exist_ok=True)

    # Template: se não enviar, usa template.docx da raiz
    template_file = request.files.get("template_file")
    if template_file and template_file.filename:
        tpl_name = secure_filename(template_file.filename)
        template_path = uploads / tpl_name
        template_file.save(template_path)
    else:
        template_path = Path("template.docx")
        if not template_path.exists():
            flash("Envie um template .docx ou coloque 'template.docx' na raiz do projeto.", "error")
            return redirect(url_for("index"))

    # Logo (opcional)
    logo_file = request.files.get("logo")
    logo_path = None
    if logo_file and logo_file.filename:
        logo_name = secure_filename(logo_file.filename)
        logo_path = uploads / logo_name
        logo_file.save(logo_path)

    # Gera o DOCX preenchido
    output = fill_docx(template_path, data, logo_path)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"PROPOSTA_{data.get('CLIENTE','cliente')}_{ts}.docx"

    return send_file(output, as_attachment=True, download_name=out_name,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
